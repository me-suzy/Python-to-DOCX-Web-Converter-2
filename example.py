from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Creăm un document nou
doc = Document()

# Adăugăm titlul
titlu = doc.add_heading('Exemplu de Document Generat', level=0)
titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adăugăm o secțiune de introducere
doc.add_paragraph('Acest document a fost generat automat folosind Python și biblioteca python-docx.')

# Adăugăm primul paragraf cu formatare
p1 = doc.add_paragraph('Acesta este primul paragraf al documentului. ')
run1 = p1.add_run('Această parte este formatată diferit - bold și colorată.')
run1.bold = True
run1.font.size = Pt(14)
run1.font.color.rgb = RGBColor(0, 0, 255)  # Albastru

# Adăugăm al doilea paragraf
p2 = doc.add_paragraph('Acesta este al doilea paragraf. ')
run2 = p2.add_run('Folosim italic și font diferit aici.')
run2.italic = True
run2.font.name = 'Arial'
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(255, 0, 0)  # Roșu

# Adăugăm o listă
doc.add_paragraph('Elemente importante:', style='List Bullet')
doc.add_paragraph('Primul element', style='List Bullet')
doc.add_paragraph('Al doilea element', style='List Bullet')
doc.add_paragraph('Al treilea element', style='List Bullet')

# Adăugăm un paragraf final
doc.add_paragraph()
final_paragraph = doc.add_paragraph('Documentul este complet!')
final_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

